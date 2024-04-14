'''
Date: 2023-10-25 15:24:45
LastEditors: Lei_T
LastEditTime: 2024-03-28 16:37:20
'''

#%%
import os
import pandas as pd 
import numpy as np
from math import floor
import docx
from docx.shared import Cm, RGBColor, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from WindPy import w
from datetime import datetime, timedelta
from openpyxl import load_workbook
from openpyxl.chart.shapes import GraphicalProperties
from openpyxl.drawing.line import LineProperties
from openpyxl.styles import numbers
from openpyxl.utils.dataframe import dataframe_to_rows
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
import matplotlib.ticker as ticker
import warnings
warnings.filterwarnings('ignore')
w.start()

def get_info(df_link,df_rating,FundCode,today,rptDate_all,language_type):
    '''
    获取基金基本信息
    '''
    if_link=type(df_link.loc[FundCode,'link_FundCode'])==list
    ###基金名称
    if language_type=='zh':
        FundName=w.wss(FundCode,"fund_info_name").Data[0][0]
    else:
        df_fundname_en=pd.read_excel(path_data+'公募基金英文名称一览表.xlsx',dtype=str, header=0, skiprows=2)
        df_fundname_en['英文名称'].fillna(method='ffill',inplace=True)
        df_fundname_en['基金代码']=df_fundname_en['基金代码'].apply(lambda x:x+'.OF')
        df_fundname_en.set_index('基金代码',inplace=True)
        FundName=df_fundname_en.loc[FundCode,'英文名称']
    ###基金类别
    FundType=w.wss(FundCode, "fund_firstinvesttype").Data[0][0]
    #投资风格
    FundStyle='--'
    if language_type=='zh':
        if FundType!='债券型基金' and FundType!='债券型基金':
            FundStyle=w.wss(FundCode, "style_marketvaluestyleattribute",f"rptDate={''.join(rptDate_all.split('-'))}").Data[0][0]
            #如果本期持仓的投资风格还未更新，则获取上一期的投资风格
            if FundStyle==None:
                if rptDate_all.split('-')[1]=='12':
                    tmp_date=f"{rptDate_all.split('-')[0]}0630"
                else:
                    tmp_date=f"{str(int(rptDate_all.split('-')[0])-1)}1231"
            FundStyle=w.wss(FundCode, "style_marketvaluestyleattribute",f"rptDate={tmp_date}").Data[0][0]
    ###基金评级
    if FundCode[:-3] not in df_rating.index:
        rating_3y='--'
        rating_5y='--'
    else:
        tmp=df_rating.loc[FundCode[:-3]]
        rating_3y=tmp.iloc[-8]
        if pd.isna(rating_3y):
            rating_3y='--'
        rating_5y=tmp.iloc[-4]
        if pd.isna(rating_5y):
            rating_5y='--'
    ###基金规模
    scale=w.wss(FundCode, "netasset_total").Data[0][0]
    if language_type=='zh':
        scale=round(scale/1e+08,2)
        scale=f'{scale}亿元'
    elif language_type=='en':
        scale=round(scale/1e+09,2)
        scale=f'{scale}'
    ###单位净值、累计净值
    if language_type=='zh':
        tmp=w.wsd(FundCode, "nav,NAV_accumulated_transform", today, today, "Days=Alldays;Fill=Previous")
    elif language_type=='en':
        tmp=w.wsd(FundCode, "nav,nav_adj", today, today, "Days=Alldays;Fill=Previous")
    tmp=pd.DataFrame(data={'time':tmp.Times,'nav1':tmp.Data[0],'nav2':tmp.Data[1]})
    tmp['time']=tmp['time'].apply(lambda x:x.strftime('%Y-%m-%d'))
    tmp.set_index('time',inplace=True)
    nav,nav2=tmp.loc[today,'nav1'],tmp.loc[today,'nav2']
    nav=str(round(nav,4))
    nav2=str(round(nav2,4))
    if nav2=='nan':
        nav2='--'
    if if_link==True and language_type=='zh' and nav2!='--':
        FundLevel=FundName[-1]
        nav=FundLevel+'类：'+nav
        nav2=FundLevel+'类：'+nav2
        for link_FundCode in df_link.loc[FundCode,'link_FundCode']:
            tmp=w.wsd(link_FundCode, "nav,NAV_accumulated_transform", today, today, "Days=Alldays;Fill=Previous")
            tmp=pd.DataFrame(data={'time':tmp.Times,tmp.Fields[0]:tmp.Data[0],tmp.Fields[1]:tmp.Data[1]})
            tmp['time']=tmp['time'].apply(lambda x:x.strftime('%Y-%m-%d'))
            tmp.set_index('time',inplace=True)
            link_nav,link_nav_acc=tmp.loc[today,'NAV'],tmp.loc[today,'NAV_ACCUMULATED_TRANSFORM']
            link_nav=str(round(link_nav,4))
            link_nav_acc=str(round(link_nav_acc,4))
            link_FundName=w.wsd(link_FundCode,"fund_info_name").Data[0][0]
            link_FundLevel=link_FundName[-1]
            link_nav=link_FundLevel+'类：'+link_nav
            link_nav_acc=link_FundLevel+'类：'+link_nav_acc
            nav=nav+'\n'+link_nav
            nav2=nav2+'\n'+link_nav_acc
    ###成立日期
    SetupDate=w.wss(FundCode,"fund_setupdate").Data[0][0]
    SetupDate=SetupDate.strftime('%Y-%m-%d')
    #如果有关联基金，则加上关联基金的成立日期
    if if_link==True and language_type=='zh':
        FundLevel=FundName[-1]
        SetupDate=FundLevel+'类：'+SetupDate
        for link_FundCode in df_link.loc[FundCode,'link_FundCode']:
            link_FundName=w.wsd(link_FundCode,"fund_info_name").Data[0][0]
            link_FundLevel=link_FundName[-1]
            link_SetupDate=w.wsd(link_FundCode,"fund_setupdate").Data[0][0]
            link_SetupDate=link_SetupDate.strftime('%Y-%m-%d')
            link_SetupDate=link_FundLevel+'类：'+link_SetupDate
            SetupDate=SetupDate+'\n'+link_SetupDate
    ###最低投资额
    MinBuy=w.wss(FundCode, "fund_pchredm_pchminamt").Data[0][0]
    if MinBuy==0.01:
        if language_type=='zh':
            MinBuy='0.01元'
        elif language_type=='en':
            MinBuy='0.01'
    else:
        if language_type=='zh':
            MinBuy=f'{int(MinBuy)}元'
        elif language_type=='en':
            MinBuy=f'{int(MinBuy)}'
    if if_link==True and language_type=='zh':
        FundLevel=FundName[-1]
        MinBuy=FundLevel+'类：'+MinBuy
        for link_FundCode in df_link.loc[FundCode,'link_FundCode']:
            link_MinBuy=w.wsd(link_FundCode,"fund_pchredm_pchminamt").Data[0][0]
            if link_MinBuy==0.01:
                link_MinBuy='0.01元'
            else:
                link_MinBuy=f'{int(link_MinBuy)}元'
            link_FundName=w.wsd(link_FundCode,"fund_info_name").Data[0][0]
            link_FundLevel=link_FundName[-1]
            link_MinBuy=link_FundLevel+'类：'+link_MinBuy
            MinBuy=MinBuy+'；'+link_MinBuy
    ###管理人
    mgrcomp=w.wsd(FundCode,"fund_mgrcomp").Data[0][0]
    ###托管人
    CustodianBank=w.wsd(FundCode,"fund_custodianbank").Data[0][0]
    ###业绩比较基准
    BenchMark=w.wsd(FundCode,"fund_benchmark").Data[0][0]
    ###投资目标
    InvestObject=w.wsd(FundCode,"fund_investobject").Data[0][0]
    ###基金经理简介
    num_managers=len(w.wss(FundCode, 'fund_fundmanageroftradedate').Data[0][0].split(','))
    resume_str=''
    for i in range(num_managers):
        name=w.wss(FundCode, 'fund_longestfundmanager_hist', f"order={i+1}").Data[0][0]
        resume=w.wss(FundCode, 'fund_manager_resume', f"order={i+1}").Data[0][0]
        if language_type=='zh':
            resume_str+=name+'，'+resume
        else:
            resume_str+=name+', '+resume
        if num_managers>1 and i!=num_managers-1:
            resume_str+='\n'
    ###收尾初修改FundName
    #之所以不在前面修改，是因为获取到FundName后，别的地方还要用到FundName，修改后别的地方会出错
    if language_type=='zh':
        if if_link:
            tmp=pd.DataFrame(columns=['level'])
            tmp.loc[FundCode.split('.')[0]]=[FundName[-1]]
            for link_FundCode in df_link.loc[FundCode,'link_FundCode']:
                link_FundName=w.wss(link_FundCode,"fund_info_name").Data[0][0]
                tmp.loc[link_FundCode.split('.')[0]]=[link_FundName[-1]]
            formatted_string = '、'.join(f"{row['level']}类：{index}" for index,row in tmp.sort_values(by='level').iterrows())
            FundName=FundName[:-1]+'（'+formatted_string+'）'#删除A基金最后的字母，加上“（A：XX、B：XX）”的字样
        else:
            FundName+='（'+FundCode.split('.')[0]+'）'
    ###写入dataframe
    df_info=pd.DataFrame(columns=['结果'])
    if if_link:
        df_info.loc['基金份额分类']=[FundLevel]
    df_info.loc['基金名称']=[FundName]
    df_info.loc['基金类别']=[FundType]
    df_info.loc['投资风格']=[FundStyle]
    df_info.loc['基金评级_三年']=[rating_3y]
    df_info.loc['基金评级_五年']=[rating_5y]
    df_info.loc['基金规模']=[scale]
    df_info.loc['单位净值']=[nav]
    if language_type=='zh':
        df_info.loc['累计净值']=[nav2]
    elif language_type=='en':
        df_info.loc['复权单位净值']=[nav2]
    df_info.loc['成立日期']=[SetupDate]
    df_info.loc['最低投资额']=[MinBuy]
    df_info.loc['基金管理人']=[mgrcomp]
    df_info.loc['托管银行']=[CustodianBank]
    df_info.loc['业绩比较基准']=[BenchMark]
    df_info.loc['投资目标']=[InvestObject]
    df_info.loc['基金经理简介']=[resume_str]
    if language_type=='en':
        df_info.loc['最低投资额','结果']=df_info.loc['最低投资额','结果'].replace('\n','；')
        df_info.rename(index={'基金类别':'Fund Category',
                              '投资风格':'Investment Style',
                              '基金规模':'Total Net Asset (date)',
                              '单位净值':'NAV (date)',
                              '复权单位净值':'Adjusted NAV (date)',
                              '成立日期':'Inception Date',
                              '最低投资额':'Min Investment(Yuan)',
                              '基金管理人':'Fund Company',
                              '托管银行':'Custodian Bank',
                              '业绩比较基准':'Benchmark',
                              '投资目标':'Fund Investment Objective',
                              '基金经理简介':'Fund Manager'},inplace=True)
    print('基本信息获取完毕')
    return df_info

def get_nav(FundCode,df_info,today,language_type):
    '''
    获取全历史净值走势
    '''
    ###指定IndexCode_1
    FundType=w.wss(FundCode, "fund_investtype").Data[0][0]
    IndexCode_2='nan'#设定一个初始值，如果是指定类型的基金，则IndexCode_2自动设置为空值，否则再自动匹配合适的指数
    if FundType=='普通股票型基金' or FundType=='Normal Equity Funds':
        IndexCode_1='885000.WI'
    elif FundType=='偏股混合型基金' or FundType=='Aggressive Allocation Funds':
        IndexCode_1='885001.WI'
    elif FundType=='平衡混合型基金' or FundType=='Balanced Funds':
        IndexCode_1='885002.WI'
    elif FundType=='偏债混合型基金' or FundType=='Moderate Allocation Funds':
        IndexCode_1='885003.WI'
        IndexCode_2=''
    elif FundType=='被动指数型基金' or FundType=='Passive Equity Index Funds':
        IndexCode_1='885004.WI'
    elif FundType=='混合债券型二级基金' or FundType=='Enhanced Bond Funds (Secondary Market)':
        IndexCode_1='885007.WI'
        IndexCode_2=''
    elif FundType=='中长期纯债型基金' or FundType=='Mid / Long-term Bond Funds':
        IndexCode_1='885008.WI'
        IndexCode_2=''
    elif FundType=='货币市场型基金' or FundType=='Money Market Funds':
        IndexCode_1='885009.WI'
        IndexCode_2=''
    elif FundType=='增强指数型基金' or FundType=='Enhanced Equity Index Funds':
        IndexCode_1='885044.WI'
    elif FundType=='灵活配置型基金' or FundType=='Flexible Allocation Funds':
        IndexCode_1='885061.WI'
    elif FundType=='短期纯债型基金' or FundType=='Short-term Bond Funds':
        IndexCode_1='885062.WI'
        IndexCode_2=''
    elif FundType=='被动指数型债券基金' or FundType=='Passive Bond Index Funds':
        IndexCode_1='885063.WI'
        IndexCode_2=''
    elif FundType=='混合型FOF基金' or FundType=='Hybrid fof fund':
        IndexCode_1='885072.WI'


    ###指定IndexCode_2
    if IndexCode_2=='nan':
        df_fund_bench=pd.read_excel(path_data+'农银基金基准.xlsx',index_col='FundCode')#之所以不从wind直接获取，是因为英文模式下不好判断IndexCode_2
        bench=df_fund_bench.loc[FundCode,'Benchmark']
        if '沪深300' in bench:
            IndexCode_2='000300.SH'
            if language_type=='zh':
                IndexCode_2_name='沪深300'
            else:
                IndexCode_2_name='CSI 300 INDEX'
        elif '中证700' in bench:
            IndexCode_2='000907.CSI'
            if language_type=='zh':
                IndexCode_2_name='中证700'
            else:
                IndexCode_2_name='CSI Small & MidCap 700 index'
        elif '中证800' in bench:
            IndexCode_2='000906.SH'
            if language_type=='zh':
                IndexCode_2_name='中证800'
            else:
                IndexCode_2_name='CSI 800 index'
        elif '中证1000' in bench:
            IndexCode_2='000852.SH'
            if language_type=='zh':
                IndexCode_2_name='中证1000'
            else:
                IndexCode_2_name='CSI 1000 Index'
        elif '国有企业综合' in bench:
            IndexCode_2='000955.CSI'
            if language_type=='zh':
                IndexCode_2_name='中证国有企业综合指数'
            else:
                IndexCode_2_name='CSI State-owned Enterprises Composite Index'
        elif '国有企业改革' in bench:
            IndexCode_2='399974.SZ'
            if language_type=='zh':
                IndexCode_2_name='中证国有企业改革指数'
            else:
                IndexCode_2_name='CSI State-Owned Enterprises Reform Index'
        elif '新能源' in bench:
            IndexCode_2='399808.SZ'
            if language_type=='zh':
                IndexCode_2_name='中证新能源指数'
            else:
                IndexCode_2_name='CSI New Energy Index'
        elif '大农业' in bench:
            IndexCode_2='399814.SZ'
            if language_type=='zh':
                IndexCode_2_name='中证大农业指数'
            else:
                IndexCode_2_name='CSI Grand Agriculture Index'
        elif '医药' in bench:
            IndexCode_2='000933.SH'
            if language_type=='zh':
                IndexCode_2_name='中证医药卫生指数'
            else:
                IndexCode_2_name='CSI Health Care Index'
        elif 'TMT' in bench:
            IndexCode_2='000998.CSI'
            if language_type=='zh':
                IndexCode_2_name='中证TMT产业主题指数'
            else:
                IndexCode_2_name='CSI TMT Industries Index'
        elif '内地消费' in bench:
            IndexCode_2='000942.CSI'
            if language_type=='zh':
                IndexCode_2_name='中证内地消费主题指数'
            else:
                IndexCode_2_name='CSI China Mainland Consumer Index'
        elif '新华社民族品牌工程' in bench:
            IndexCode_2='931403.CSI'
            if language_type=='zh':
                IndexCode_2_name='中证新华社民族品牌工程指数'
            else:
                IndexCode_2_name='CSI National Brands Project of Xinhua Index'
        elif '新兴产业' in bench:
            IndexCode_2='000964.CSI'
            if language_type=='zh':
                IndexCode_2_name='中证新兴产业指数'
            else:
                IndexCode_2_name='CSI Emerging Industries index'
        else:
            IndexCode_2=''

    ###基金净值
    df_nav_1=pd.read_excel(path_data+'农银基金净值.xlsx',sheet_name='净值')
    df_nav_1.rename(columns={df_nav_1.columns[0]:'TradingDay'},inplace=True)
    df_nav_1['TradingDay']=df_nav_1['TradingDay'].apply(pd.Timestamp)
    df_nav_1=df_nav_1.loc[:,['TradingDay',FundCode]]

    valid_index=(df_nav_1[df_nav_1[FundCode]>0]).index
    df_nav_1=df_nav_1.loc[valid_index]
    
    FundName=df_info.loc['基金名称','结果'].split('（')[0]
    if language_type=='zh':#如果是中文，则图例末尾加上“净值”字样
        FundName+='净值'
    elif language_type=='en':#如果是英文，则不加
        pass
    df_nav_1.rename(columns={FundCode:FundName},inplace=True)    

    ###基准净值
    df_nav_2=pd.read_excel(path_data+'农银基金净值.xlsx',sheet_name='基准指数价格')
    df_nav_2.rename(columns={df_nav_2.columns[0]:'TradingDay'},inplace=True)
    df_nav_2['TradingDay']=df_nav_2['TradingDay'].apply(pd.Timestamp)
    BchCode=FundCode.split('.')[0]+'BI.WI'
    df_nav_2=df_nav_2.loc[valid_index,['TradingDay',BchCode]]

    if language_type=='zh':
        BchName='业绩比较基准'
    elif language_type=='en':
        BchName='Benchmark'
    df_nav_2.rename(columns={BchCode:BchName},inplace=True)

    ###指数1净值
    df_nav_3=pd.read_excel(path_data+'农银基金净值.xlsx',sheet_name='基金指数价格')
    df_nav_3.rename(columns={df_nav_3.columns[0]:'TradingDay'},inplace=True)
    df_nav_3['TradingDay']=df_nav_3['TradingDay'].apply(pd.Timestamp)
    df_nav_3=df_nav_3.loc[valid_index,['TradingDay',IndexCode_1]]

    IndexName_1=w.wss(IndexCode_1, "sec_name").Data[0][0]
    if language_type=='zh':#如果是中文，则图例末尾加上“净值”字样
        if IndexCode_1=='885008.WI':
            IndexName_1='中长期纯债型基金指数净值'
        else:
            IndexName_1+='净值'
    elif language_type=='en':#如果是英文，则不加
        if IndexName_1.split(' ')[0]=='wind':
            IndexName_1=' '.join(IndexName_1.split(' ')[1:])
    df_nav_3.rename(columns={IndexCode_1:IndexName_1},inplace=True)

    ###指数2净值
    if IndexCode_2!='':
        df_nav_4=pd.read_excel(path_data+'农银基金净值.xlsx',sheet_name='宽基指数价格')
        df_nav_4.rename(columns={df_nav_4.columns[0]:'TradingDay'},inplace=True)
        df_nav_4['TradingDay']=df_nav_4['TradingDay'].apply(pd.Timestamp)
        df_nav_4=df_nav_4.loc[valid_index,['TradingDay',IndexCode_2]]
        df_nav_4.rename(columns={IndexCode_2:IndexCode_2_name},inplace=True)

    ###合并
    df_price=df_nav_1.merge(df_nav_2,on='TradingDay',how='left')
    df_price=df_price.merge(df_nav_3,on='TradingDay',how='left')
    if IndexCode_2!='':
        df_price=df_price.merge(df_nav_4,on='TradingDay',how='left')
    df_price.fillna(method='ffill',inplace=True)
    df_price.dropna(axis=1,how='all',inplace=True)
    df_price.dropna(axis=0,inplace=True)#防止开头基金净值为一直为1、基准没有数据的情况

    ###归一化
    factor=df_price[FundName].iloc[0]
    if BchName in df_price.columns:
        df_price[BchName]/=df_price[BchName].iloc[0]
        df_price[BchName]*=factor
    df_price[IndexName_1]/=df_price[IndexName_1].iloc[0]
    df_price[IndexName_1]*=factor
    if IndexCode_2!='':
        df_price[IndexCode_2_name]/=df_price[IndexCode_2_name].iloc[0]
        df_price[IndexCode_2_name]*=factor   
    df_price.set_index('TradingDay',inplace=True)
    df_price.dropna(axis=1,how='all',inplace=True)#000259.OF没有基准数据，需要drop

    print('净值走势获取完毕')
    return df_price

def get_asset(FundCode,rptDate_top,language_type):
    ###资产配置
    ratio_Stock=w.wss(FundCode, "prt_stocktoasset",f"rptDate={rptDate_top}").Data[0][0]
    ratio_Bond=w.wss(FundCode, "prt_bondtoasset",f"rptDate={rptDate_top}").Data[0][0]
    ratio_Fund=w.wss(FundCode, "prt_fundtoasset",f"rptDate={rptDate_top}").Data[0][0]
    ratio_Warrant=w.wss(FundCode, "prt_warranttoasset",f"rptDate={rptDate_top}").Data[0][0]
    ratio_Cash=w.wss(FundCode, "prt_cashtoasset",f"rptDate={rptDate_top}").Data[0][0]
    ratio_Other=w.wss(FundCode, "prt_othertoasset",f"rptDate={rptDate_top}").Data[0][0]
    ratio_Sec=w.wss(FundCode, "prt_seclendingvaluetoasset",f"rptDate={rptDate_top}").Data[0][0]
    ratio_mmf=w.wss(FundCode, "mmf_reverserepotoasset",f"rptDate={rptDate_top}").Data[0][0]

    df_asset=pd.DataFrame(columns=['占总资产比例'])
    df_asset.loc['股票']=[ratio_Stock]
    df_asset.loc['债券']=[ratio_Bond]
    df_asset.loc['基金']=[ratio_Fund]
    df_asset.loc['权证']=[ratio_Warrant]
    df_asset.loc['银行存款']=[ratio_Cash]
    df_asset.loc['其他资产']=[ratio_Other]
    df_asset.loc['转融通证券出借业务']=[ratio_Sec]
    df_asset.loc['买入返售证券']=[ratio_mmf]
    df_asset=df_asset/100
    if language_type=='en':
        df_asset.fillna(0,inplace=True)
        df_asset.loc['Stock']=[df_asset.loc['股票','占总资产比例']]
        df_asset.loc['Bond']=[df_asset.loc['债券','占总资产比例']]
        df_asset.loc['Cash']=[df_asset.loc['银行存款','占总资产比例']]
        df_asset.loc['Others']=[df_asset.loc['基金','占总资产比例']+
                                df_asset.loc['权证','占总资产比例']+
                                df_asset.loc['其他资产','占总资产比例']+
                                df_asset.loc['转融通证券出借业务','占总资产比例']+
                                df_asset.loc['买入返售证券','占总资产比例']]
        df_asset=df_asset.loc[['Stock','Bond','Cash','Others']]
    df_asset.dropna(inplace=True)
    df_asset.sort_values(by='占总资产比例',ascending=False,inplace=True)
    print('资产配置获取完毕')
    return df_asset


def get_top(FundCode,rptDate_top):
    df_top=pd.DataFrame(columns=['股票名称','股票占净值比','债券名称','债券占净值比'],index=range(1,11))
    #获取股票前十大持仓
    for n in range(1,11):
        tmp=w.wss(FundCode,'prt_topstockname',f'rptDate={rptDate_top};order={n}').Data[0][0]
        if pd.isna(tmp)==False:
            df_top.loc[n,'股票名称']=tmp
            tmp_weight=f'{w.wss(FundCode, "prt_heavilyheldstocktonav",f"rptDate={rptDate_top};order={n}").Data[0][0]/100:.2%}'
            df_top.loc[n,'股票占净值比']=tmp_weight
        else:
            break
    #获取债券前五大持仓
    for n in range(1,6):
        tmp=w.wss(FundCode,'prt_topbondname',f'rptDate={rptDate_top};order={n}').Data[0][0]
        if pd.isna(tmp)==False:
            df_top.loc[n,'债券名称']=tmp
            tmp_weight=f'{w.wss(FundCode, "prt_heavilyheldbondtonav",f"rptDate={rptDate_top};order={n}").Data[0][0]/100:.2%}'
            df_top.loc[n,'债券占净值比']=tmp_weight
        else:
            break
    df_top.fillna('--',inplace=True)
    print('前n大配置获取完毕')
    return df_top


def get_rank(df_FundType,FundCode,today):
    ###筛选出同类型基金列表
    FundInvestType=df_FundType.loc[FundCode,'三级分类']
    df_FundType_part=df_FundType[df_FundType['三级分类']==FundInvestType]
    ###同类型基金列表中，筛选出尚未到期的基金
    df_FundType_part=df_FundType_part.loc[df_FundType_part.index.isin(df_fund_special_date['FundCode'].unique())]
    str_funds=','.join(df_FundType_part.index.tolist())


    ###获取各时间点
    #上年年末时间点
    time_ytd=datetime(datetime.strptime(today, '%Y-%m-%d').year-1 , 12, 31).strftime('%Y-%m-%d')
    #前3月、6月、1年、2年时间点
    time_3m=w.tdaysoffset(-3, today, "Days=Alldays;Period=M").Data[0][0].strftime('%Y-%m-%d')
    time_6m=w.tdaysoffset(-6, today, "Days=Alldays;Period=M").Data[0][0].strftime('%Y-%m-%d')
    time_1y=w.tdaysoffset(-1, today, "Days=Alldays;Period=Y").Data[0][0].strftime('%Y-%m-%d')
    time_2y=w.tdaysoffset(-2, today, "Days=Alldays;Period=Y").Data[0][0].strftime('%Y-%m-%d')
    #成立的时间点
    #由于成立日的净值可能为nan值，故需要寻找净值披露首日
    se_fund_special_date_part=df_fund_special_date[df_fund_special_date['FundCode']==FundCode].iloc[0]
    time_setup=se_fund_special_date_part['净值披露首日'].strftime('%Y-%m-%d')


    ###检测是否要创建文件或更新数据
    flag1=True#是否存有截面净值文件
    flag2=True#截面净值文件是否有指定类型的数据
    if os.path.exists(path_data+f'截面净值_{today}.pkl'):
        df=pd.read_pickle(path_data+f'截面净值_{today}.pkl')
        if FundInvestType not in df['三级分类'].unique():
            flag2=False
    else:
        flag1=False

    if flag1==True and flag2==True:
        print('已存在数据')
        if time_setup in df.columns:
            print('数据已存在成立日净值，不再更新')
        else:
            print('数据无存在成立日净值，进行更新')
            tmp=w.wss(str_funds, "NAV_adjusted_transform",f"tradeDate={time_setup}")
            tmp=pd.DataFrame({'FundCode': tmp.Codes,time_setup:tmp.Data[0]})
            tmp.set_index('FundCode',inplace=True)
            df=df.merge(tmp,left_index=True,right_index=True,how='outer')
            df.to_pickle(path_data+f'截面净值_{today}.pkl')
    else:
        print('不存在数据，开始更新')
        ###获取各时间点的净值
        dic={'今日':today,time_setup:time_setup,
             '年初':time_ytd,'三月前':time_3m,
             '六月前':time_6m,'一年前':time_1y,
             '两年前':time_2y}
        lst_df=[]
        for key in dic.keys():
            time=dic[key]
            tmp=w.wss(str_funds, "NAV_adjusted_transform",f"tradeDate={time}")
            tmp=pd.DataFrame({'FundCode': tmp.Codes,key:tmp.Data[0]})
            tmp.set_index('FundCode',inplace=True)
            lst_df.append(tmp)
        df_new=pd.concat(lst_df,axis=1,join='outer')
        df_new['三级分类']=FundInvestType

        ###合并到已有数据，并保存
        if flag1==False:
            df=df_new.copy()
        elif flag2==False:
            df=pd.concat([df,df_new])
        df.to_pickle(path_data+f'截面净值_{today}.pkl')
        
    ####在更新好数据的基础上，筛选出同类基金的截面净值数据
    df=df[df['三级分类']==FundInvestType]
    ###计算各收益率
    df['收益_成立以来']=df['今日']/df[time_setup]-1
    df['收益_ytd']=df['今日']/df['年初']-1
    df['收益_3m']=df['今日']/df['三月前']-1
    df['收益_6m']=df['今日']/df['六月前']-1
    df['收益_1y']=df['今日']/df['一年前']-1
    df['收益_2y']=df['今日']/df['两年前']-1

    #如果成立日期比今年年初/前三月...大，则不计算收益率和排名
    ret_setup_self=f"{df.loc[FundCode,'收益_成立以来']:.2%}"
    if datetime.strptime(time_setup,'%Y-%m-%d')>datetime.strptime(time_ytd,'%Y-%m-%d'):
        ret_ytd_self='--'
    else:
        ret_ytd_self=f"{df.loc[FundCode,'收益_ytd']:.2%}"
    if datetime.strptime(time_setup,'%Y-%m-%d')>datetime.strptime(time_3m,'%Y-%m-%d'):
        ret_3m_self='--'
    else:
        ret_3m_self=f"{df.loc[FundCode,'收益_3m']:.2%}"
    if datetime.strptime(time_setup,'%Y-%m-%d')>datetime.strptime(time_6m,'%Y-%m-%d'):
        ret_6m_self='--'
    else:
        ret_6m_self=f"{df.loc[FundCode,'收益_6m']:.2%}"
    if datetime.strptime(time_setup,'%Y-%m-%d')>datetime.strptime(time_1y,'%Y-%m-%d'):
        ret_1y_self='--'
    else:
        ret_1y_self=f"{df.loc[FundCode,'收益_1y']:.2%}"
    if datetime.strptime(time_setup,'%Y-%m-%d')>datetime.strptime(time_2y,'%Y-%m-%d'):
        ret_2y_self='--'
    else:
        ret_2y_self=f"{df.loc[FundCode,'收益_2y']:.2%}"

    ###计算排名
    ret_setup_rank=f"{int(df['收益_成立以来'].dropna().rank(ascending=False,na_option='bottom').loc[FundCode])}/{int(len(df['收益_成立以来'].dropna()))}"
    if datetime.strptime(time_setup,'%Y-%m-%d')>datetime.strptime(time_ytd,'%Y-%m-%d'):
        ret_ytd_rank='--'
    else:
        ret_ytd_rank=f"{int(df['收益_ytd'].dropna().rank(ascending=False,na_option='bottom').loc[FundCode])}/{int(len(df['收益_ytd'].dropna()))}"
    if datetime.strptime(time_setup,'%Y-%m-%d')>datetime.strptime(time_3m,'%Y-%m-%d'):
        ret_3m_rank='--'
    else:
        ret_3m_rank=f"{int(df['收益_3m'].dropna().rank(ascending=False,na_option='bottom').loc[FundCode])}/{int(len(df['收益_3m'].dropna()))}"
    if datetime.strptime(time_setup,'%Y-%m-%d')>datetime.strptime(time_6m,'%Y-%m-%d'):
        ret_6m_rank='--'
    else:
        ret_6m_rank=f"{int(df['收益_6m'].dropna().rank(ascending=False,na_option='bottom').loc[FundCode])}/{int(len(df['收益_6m'].dropna()))}"
    if datetime.strptime(time_setup,'%Y-%m-%d')>datetime.strptime(time_1y,'%Y-%m-%d'):
        ret_1y_rank='--'
    else:
        ret_1y_rank=f"{int(df['收益_1y'].dropna().rank(ascending=False,na_option='bottom').loc[FundCode])}/{int(len(df['收益_1y'].dropna()))}"
    if datetime.strptime(time_setup,'%Y-%m-%d')>datetime.strptime(time_2y,'%Y-%m-%d'):
        ret_2y_rank='--'
    else:
        ret_2y_rank=f"{int(df['收益_2y'].dropna().rank(ascending=False,na_option='bottom').loc[FundCode])}/{int(len(df['收益_2y'].dropna()))}"


    ###计算同类平均收益率
    ret_setup_avg=f"{df['收益_成立以来'].dropna().mean():.2%}"
    ret_ytd_avg=f"{df['收益_ytd'].dropna().mean():.2%}"
    ret_3m_avg=f"{df['收益_3m'].dropna().mean():.2%}"
    ret_6m_avg=f"{df['收益_6m'].dropna().mean():.2%}"
    ret_1y_avg=f"{df['收益_1y'].dropna().mean():.2%}"
    ret_2y_avg=f"{df['收益_2y'].dropna().mean():.2%}"


    ###计算业绩基准收益率
    BchCode=w.wss(FundCode, 'fund_benchindexcode').Data[0][0]
    ret_setup_bch=f'{(((w.wss(BchCode, "close",f"tradeDate={today}").Data[0][0]/w.wss(BchCode, "close",f"tradeDate={time_setup}").Data[0][0]))-1):.2%}'
    if datetime.strptime(time_setup,'%Y-%m-%d')>datetime.strptime(time_ytd,'%Y-%m-%d'):
        ret_ytd_bch='--'
    else:
        ret_ytd_bch=f'{(((w.wss(BchCode, "close",f"tradeDate={today}").Data[0][0]/w.wss(BchCode, "close",f"tradeDate={time_ytd}").Data[0][0]))-1):.2%}'
    if datetime.strptime(time_setup,'%Y-%m-%d')>datetime.strptime(time_3m,'%Y-%m-%d'):
        ret_3m_bch='--'
    else:
        ret_3m_bch=f'{(((w.wss(BchCode, "close",f"tradeDate={today}").Data[0][0]/w.wss(BchCode, "close",f"tradeDate={time_3m}").Data[0][0]))-1):.2%}'
    if datetime.strptime(time_setup,'%Y-%m-%d')>datetime.strptime(time_6m,'%Y-%m-%d'):
        ret_6m_bch='--'
    else:
        ret_6m_bch=f'{(((w.wss(BchCode, "close",f"tradeDate={today}").Data[0][0]/w.wss(BchCode, "close",f"tradeDate={time_6m}").Data[0][0]))-1):.2%}'
    if datetime.strptime(time_setup,'%Y-%m-%d')>datetime.strptime(time_1y,'%Y-%m-%d'):
        ret_1y_bch='--'
    else:
        ret_1y_bch=f'{(((w.wss(BchCode, "close",f"tradeDate={today}").Data[0][0]/w.wss(BchCode, "close",f"tradeDate={time_1y}").Data[0][0]))-1):.2%}'
    if datetime.strptime(time_setup,'%Y-%m-%d')>datetime.strptime(time_2y,'%Y-%m-%d'):
        ret_2y_bch='--'
    else:
        ret_2y_bch=f'{(((w.wss(BchCode, "close",f"tradeDate={today}").Data[0][0]/w.wss(BchCode, "close",f"tradeDate={time_2y}").Data[0][0]))-1):.2%}'


    ###存入dataframe
    df_rank=pd.DataFrame(index=['今年以来','近三个月','近六个月',
                                '近一年','近两年','成立以来'],
                         columns=['本基金收益','同类排名','同类平均','比较基准'])
    df_rank.iloc[0]=[ret_ytd_self,ret_ytd_rank,ret_ytd_avg,ret_ytd_bch]
    df_rank.iloc[1]=[ret_3m_self,ret_3m_rank,ret_3m_avg,ret_3m_bch]
    df_rank.iloc[2]=[ret_6m_self,ret_6m_rank,ret_6m_avg,ret_6m_bch]
    df_rank.iloc[3]=[ret_1y_self,ret_1y_rank,ret_1y_avg,ret_1y_bch]
    df_rank.iloc[4]=[ret_2y_self,ret_2y_rank,ret_2y_avg,ret_2y_bch]
    df_rank.iloc[5]=[ret_setup_self,ret_setup_rank,ret_setup_avg,ret_setup_bch]

    df_rank=df_rank.T
    df_rank.loc['']=['','','','','','']
    df_rank.loc['截止日期']=[today,'','','','','']
    df_rank.loc['年初日期']=[time_ytd,'','','','','']
    df_rank.loc['前三个月日期']=[time_3m,'','','','','']
    df_rank.loc['前六个月日期']=[time_6m,'','','','','']
    df_rank.loc['前一年日期']=[time_1y,'','','','','']
    df_rank.loc['前两年日期']=[time_2y,'','','','','']
    df_rank.loc[f'{FundCode}成立日期']=[time_setup,'','','','','']
    if language_type=='en':
        df_rank.rename(index={'本基金收益':'Total Return',
                              '同类排名':'Rank',
                              '同类平均':'Category Mean',
                              '比较基准':'benchmark'},inplace=True)
    print('业绩排名获取完毕')
    return df_rank



def get_fee(df_link,FundCode,language_type):
    ###申购费
    text=w.wss(FundCode, "fund_purchasefee","chargesType=0").Data[0][0]
    df_buyfee=pd.DataFrame(columns=['group','day','fee'])
    for i in text.split(';\r\n'):
        if language_type=='en':
            if re.search(r'common group',i,re.IGNORECASE):
                #text拆分成群体、投资门槛、费率
                group,thresh_and_fee=i.split(':')
                if len(thresh_and_fee.split(' '))>=2:
                    thresh=thresh_and_fee.split(' ')[0]
                    fee=thresh_and_fee.split(thresh+' ')[1]
                else:
                    thresh=''
                    fee=thresh_and_fee
                #处理诸如英文版本下投资门槛返回值为“100万CNY以下”的问题
                if thresh!='':
                    digtals=re.findall(r'\d+', thresh)
                    if len(digtals)==1:
                        transfer_digtal=round(float(digtals[0])/100,5)#取四舍五入，是为了防止某个整数除以100后，出现x.00000001这种情况
                        #transfer_digtal可能是x.x、x.0两种情况，故需要再进行加工
                        if floor(transfer_digtal) == transfer_digtal:
                            transfer_digtal=int(transfer_digtal) 
                        if '以下' in thresh:
                            thresh=f'{transfer_digtal} million CNY below'
                        elif '以上' in thresh:
                            thresh=f'{transfer_digtal} million CNY above'
                    elif len(digtals)==2:
                        transfer_digtal_1=round(float(digtals[0])/100,5)
                        if floor(transfer_digtal_1) == transfer_digtal_1:
                            transfer_digtal_1=int(transfer_digtal_1) 
                        transfer_digtal_2=round(float(digtals[1])/100,5)
                        if floor(transfer_digtal_2) == transfer_digtal_2:
                            transfer_digtal_2=int(transfer_digtal_2)                  
                        thresh=f'{transfer_digtal_1}~{transfer_digtal_2} million CNY'
                df_buyfee.loc[len(df_buyfee)]=[group,thresh,fee]
        elif language_type=='zh':
            if re.search(r'普通投资群体',i,re.IGNORECASE):
                #text拆分成群体、投资门槛、费率
                group,thresh_and_fee=i.split(':')
                if len(thresh_and_fee.split(' '))>=2:
                    thresh=thresh_and_fee.split(' ')[0]
                    fee=thresh_and_fee.split(thresh+' ')[1]
                else:
                    thresh=''
                    fee=thresh_and_fee
                df_buyfee.loc[len(df_buyfee)]=[group,thresh,fee]
    #按费率排序
    def custom_sort(value):
        if value[-1]!='%':
            return -1#x元/笔等字样排在最后
        else:
            return float(value[:-1])#返回删除百分号后的数值，函数外从大到小排序
    df_buyfee=df_buyfee.sort_values(by='fee', key=lambda x: x.map(custom_sort),ascending=False)
    #添加新列
    df_buyfee['费用大类']='前端申购费率'
    df_buyfee['fee_type']=df_buyfee.index.map(lambda x:f'申购费{x+1}')
    df_buyfee.set_index('fee_type',inplace=True)
    #加上联接基金
    if_link=type(df_link.loc[FundCode,'link_FundCode'])==list
    if if_link and language_type=='zh':
        FundLevel=w.wsd(FundCode,'fund_info_name').Data[0][0][-1]
        df_buyfee['基金类别']=FundLevel
        for link_FundCode in df_link.loc[FundCode,'link_FundCode']:
            df_buyfee2=pd.DataFrame(columns=['group','day','fee','基金类别'])
            text=w.wss(link_FundCode, "fund_purchasefee","chargesType=0").Data[0][0]
            link_FundLevel=w.wsd(link_FundCode,'fund_info_name').Data[0][0][-1]
            for i in text.split(';\r\n'):
                if re.search(r'普通投资群体',i,re.IGNORECASE):
                    #text拆分成群体、投资门槛、费率
                    group,thresh_and_fee=i.split(':')
                    if len(thresh_and_fee.split(' '))>=2:
                        thresh=thresh_and_fee.split(' ')[0]
                        fee=thresh_and_fee.split(thresh+' ')[1]
                    else:
                        thresh=''
                        fee=thresh_and_fee
                    df_buyfee2.loc[len(df_buyfee2)]=[group,thresh,fee,link_FundLevel]
            df_buyfee2['费用大类']='前端申购费率'
            df_buyfee2['fee_type']=df_buyfee2.index.map(lambda x:f'申购费{x+1}')
            df_buyfee2.set_index('fee_type',inplace=True)
            df_buyfee=pd.concat([df_buyfee,df_buyfee2])



    ###赎回费
    text=w.wss(FundCode, "fund_redemptionfee","chargesType=1").Data[0][0]
    df_sellfee=pd.DataFrame(columns=['group','day','fee'])
    for i in text.split(';\r\n'):
        if language_type=='en':
            if re.search(r'ordinary investment group',i,re.IGNORECASE):
                #text拆分成群体、投资门槛、费率
                group,thresh_and_fee=i.split(':')
                if len(thresh_and_fee.split(' '))>=2:
                    thresh=thresh_and_fee.split(' ')[0]
                    fee=thresh_and_fee.split(thresh+' ')[1]
                else:
                    thresh=''
                    fee=thresh_and_fee
                #处理诸如英文版本下投资门槛返回值为“100万CNY以下”的问题
                if thresh!='':
                    digtals=re.findall(r'\d+', thresh)
                    if len(digtals)==1:
                        if '以下' in thresh.lower() or 'below' in thresh.lower():
                            if 'day' in thresh.lower():
                                thresh=f'{digtals[0]} day below'
                            elif 'year' in thresh.lower():
                                thresh=f'{digtals[0]} year below'
                        elif '以上' in thresh.lower() or 'above' in thresh.lower():
                            if 'day' in thresh.lower():
                                thresh=f'{digtals[0]} day above'
                            elif 'year' in thresh.lower():
                                thresh=f'{digtals[0]} year above'
                    elif len(digtals)==2:
                        if 'day' in thresh.lower():
                            thresh=f'{digtals[0]}~{digtals[1]} day'
                        elif 'year' in thresh.lower():
                            thresh=f'{digtals[0]}~{digtals[1]} year'
                df_sellfee.loc[len(df_sellfee)]=[group,thresh,fee]
        elif language_type=='zh':
            if re.search(r'普通投资群体',i,re.IGNORECASE):
                #text拆分成群体、投资门槛、费率
                group,thresh_and_fee=i.split(':')
                if len(thresh_and_fee.split(' '))>=2:
                    thresh=thresh_and_fee.split(' ')[0]
                    fee=thresh_and_fee.split(thresh+' ')[1]
                else:
                    thresh=''
                    fee=thresh_and_fee
                df_sellfee.loc[len(df_sellfee)]=[group,thresh,fee]
    df_sellfee['fee_type']='赎回费'
    #按费率排序
    df_sellfee=df_sellfee.sort_values(by='fee', key=lambda x: x.map(custom_sort),ascending=False)
    #添加新列
    df_sellfee['费用大类']='赎回费率'
    df_sellfee['fee_type']=df_sellfee.index.map(lambda x:f'赎回费{x+1}')
    df_sellfee.set_index('fee_type',inplace=True)
    #加上联接基金
    if_link=type(df_link.loc[FundCode,'link_FundCode'])==list
    if if_link and language_type=='zh':
        FundLevel=w.wsd(FundCode,'fund_info_name').Data[0][0][-1]
        df_sellfee['基金类别']=FundLevel
        for link_FundCode in df_link.loc[FundCode,'link_FundCode']:
            df_sellfee2=pd.DataFrame(columns=['group','day','fee','基金类别'])
            text=w.wss(link_FundCode, "fund_redemptionfee","chargesType=1").Data[0][0]
            link_FundLevel=w.wsd(link_FundCode,'fund_info_name').Data[0][0][-1]
            for i in text.split(';\r\n'):
                if re.search(r'普通投资群体',i,re.IGNORECASE):
                    #text拆分成群体、投资门槛、费率
                    group,thresh_and_fee=i.split(':')
                    if len(thresh_and_fee.split(' '))>=2:
                        thresh=thresh_and_fee.split(' ')[0]
                        fee=thresh_and_fee.split(thresh+' ')[1]
                    else:
                        thresh=''
                        fee=thresh_and_fee
                    df_sellfee2.loc[len(df_sellfee2)]=[group,thresh,fee,link_FundLevel]
            df_sellfee2['费用大类']='赎回费率'
            df_sellfee2['fee_type']=df_sellfee2.index.map(lambda x:f'赎回费{x+1}')
            df_sellfee2.set_index('fee_type',inplace=True)
            df_sellfee=pd.concat([df_sellfee,df_sellfee2])



    ###管理费、托管费、服务费
    ManagementFee,CustodianFee,SaleFee=[x for sublist in w.wss(FundCode, "fund_managementfeeratio,fund_custodianfeeratio,fund_salefeeratio", "chargesType=0").Data for x in sublist]
    if pd.isna(ManagementFee):
        ManagementFee='--%'
    else:
        ManagementFee=f'{ManagementFee/100:.2%}'
    if pd.isna(CustodianFee):
        CustodianFee='--%'
    else:
        CustodianFee=f'{CustodianFee/100:.2%}'
    if pd.isna(SaleFee):
        SaleFee='--%'
    else:
        SaleFee=f'{SaleFee/100:.2%}'
    df_otherfee=pd.DataFrame(data=[['管理费',ManagementFee],
                                   ['托管费',CustodianFee],
                                   ['销售费',SaleFee]],
                             columns=['fee_type','fee'])
    df_otherfee['费用大类']='费用信息'
    df_otherfee.set_index('fee_type',inplace=True)
    #加上联接基金
    if_link=type(df_link.loc[FundCode,'link_FundCode'])==list
    if if_link and language_type=='zh':
        FundLevel=w.wsd(FundCode,'fund_info_name').Data[0][0][-1]
        df_otherfee['基金类别']=FundLevel
        for link_FundCode in df_link.loc[FundCode,'link_FundCode']:
            ManagementFee,CustodianFee,SaleFee=[x for sublist in w.wss(link_FundCode, "fund_managementfeeratio,fund_custodianfeeratio,fund_salefeeratio", "chargesType=0").Data for x in sublist]
            if pd.isna(ManagementFee):
                ManagementFee='--%'
            else:
                ManagementFee=f'{ManagementFee/100:.2%}'
            if pd.isna(CustodianFee):
                CustodianFee='--%'
            else:
                CustodianFee=f'{CustodianFee/100:.2%}'
            if pd.isna(SaleFee):
                SaleFee='--%'
            else:
                SaleFee=f'{SaleFee/100:.2%}'
            df_otherfee2=pd.DataFrame(data=[['管理费',ManagementFee],
                                           ['托管费',CustodianFee],
                                           ['销售费',SaleFee]],
                                     columns=['fee_type','fee'])
            df_otherfee2['费用大类']='费用信息'
            FundLevel=w.wsd(link_FundCode,'fund_info_name').Data[0][0][-1]
            df_otherfee2['基金类别']=FundLevel
            df_otherfee2.set_index('fee_type',inplace=True)
            df_otherfee=pd.concat([df_otherfee,df_otherfee2])

    ###合并费率
    df_feeInfo=pd.concat([df_buyfee,df_sellfee,df_otherfee])
    if language_type=='zh':
        df_feeInfo.rename(index={'销售费':'销售服务费'},inplace=True)
    else:
        df_feeInfo['day'].fillna('--',inplace=True)
        df_feeInfo.rename(index={'管理费':'Management',
                                 '托管费':'Custodian',
                                 '销售费':'Sales&Service'},inplace=True)
    df_feeInfo['fee']=df_feeInfo['fee'].replace('--%','--')
    print('费率信息获取完毕')
    return df_feeInfo




def save(FundCode,df_info,df_price,df_asset,df_top,df_rank,df_fee):
    wb = load_workbook(path_data+'净值图表模板.xlsx')

    ###写入df_asset
    ws = wb['资产配置']
    n=2
    for i in range(len(df_asset.index)):
        ws[f'A{n}'] = df_asset.index[i]
        n+=1
    q=0
    name=df_asset.columns[q]
    ws[f'{chr(66+q)}1'] = name
    n=2
    for i in range(len(df_asset.index)):
        ws[f'{chr(66+q)}{n}'] = df_asset[name].iloc[i]
        ws[f'{chr(66+q)}{n}'].number_format = '0.00%'
        n+=1

    ###写入df_price
    ws = wb['净值走势']
    n=2
    for i in range(len(df_price.index)):
        ws[f'A{n}'] = pd.Timestamp(df_price.index[i])
        ws[f'A{n}'].number_format = numbers.FORMAT_DATE_YYYYMMDD2
        n+=1
    for q in range(len(df_price.columns)):
        name=df_price.columns[q]
        ws[f'{chr(66+q)}1'] = name
        n=2
        for i in range(len(df_price.index)):
            ws[f'{chr(66+q)}{n}'] = df_price[name].iloc[i]
            n+=1
    
    ###写入df_info
    worksheet = wb.create_sheet(title='基金基本信息')
    for row in dataframe_to_rows(df_info, index=True, header=True):
        worksheet.append(row)

    ###写入df_fee
    worksheet = wb.create_sheet(title='费率信息')
    for row in dataframe_to_rows(df_fee, index=True, header=True):
        worksheet.append(row)

    ###写入df_rank
    worksheet = wb.create_sheet(title='业绩排名')
    for row in dataframe_to_rows(df_rank, index=True, header=True):
        worksheet.append(row)
    
    ###写入df_top
    worksheet = wb.create_sheet(title='前n大股票债券')
    for row in dataframe_to_rows(df_top, index=True, header=True):
        worksheet.append(row)
    
    wb.save(path_result+f'{FundCode}结果汇总.xlsx')
    print('数据写入完毕')




###############填充word基金信息表格###############
def show_tableContent(doc,table,filter=True):
    """
    Description
    ----------
    展示word表格中所有内容

    Parameters
    ----------
    doc: docx.document.Document 读取的word主体
    table: int. 需要展示的表格位置

    """
    tables = doc.tables    #获取文件中的表格集
    table = tables[table]  
    empty = True
    info = []
    if language_type=='zh':
        columns_name = ['产品亮点','投资目标','基金经理简介','投资特色']
    else:
        columns_name = ['Fund Company','Custodian Bank','Benchmark','Fund Investment Objective','Fund Manager']
    columns_names_backup = columns_name.copy()
    rev_table_columns = []
    if language_type=='zh':
        rev_table_columns_name = ['今年以来','近三个月','近六个月','近一年','近两年','成立以来']
    else:
        rev_table_columns_name = ['YTD','3 Months','6 Months','1 Year','2 Years','Since Inception']
    for i in range(0,len(table.rows)):  #从表格第二行开始循环读取表格数据
        for j in range(0,len(table.columns)):
            try:
                # print([i,j])
                # print(table.cell(i,j).text)
                if not empty and table.cell(i,j).text == '' and content not in columns_names_backup:
                    content_loc = [content,[i,j]]
                    empty = True
                    info.append(content_loc)
                elif table.cell(i,j).text == '':
                    empty = True
                else:
                    if table.cell(i,j).text in columns_name:
                        content_loc = [table.cell(i,j).text,[i+1,j]]
                        info.append(content_loc)
                        columns_name.remove(table.cell(i,j).text)
                    if table.cell(i,j).text in rev_table_columns_name:
                        rev_table_columns.append(j)
                        rev_table_columns_name.remove(table.cell(i,j).text)
                    empty = False    
                content = table.cell(i,j).text
                if language_type=='zh':
                    if content in ['三年','五年']:
                        content = '基金评级_'+content
                    if filter:
                        content = re.search(r'[a-zA-Z0-9_\u4e00-\u9fa5\-]+',content.replace('\n','')).group()
                        
            except:
                continue
    for i in info:
        if language_type=='zh':
            if i[0] in ['本基金收益','同类排名','同类平均','比较基准']:
                del i[1][-1]
                i[1].append(rev_table_columns)
        else:
            if i[0] in ['Total Return','Rank','Category Mean','benchmark']:
                del i[1][-1]
                i[1].append(rev_table_columns)
    print('文档位置信息已保存')
    return info
            


def tableText_replace(doc,table,row,col,replacement,font='微软雅黑',C_font='微软雅黑',font_size=8,font_color='black',bold=False,center=False):
    """
    Description
    ----------
    替换word表格中内容
    
    Parameters
    ----------
    doc: docx.document.Document 读取的word主体
    table: int. 需要修改的表格位置
    row: int. 需要替换的内容所处行数
    col: int. 需要替换的内容所处列数
    replacement: str. 所替换的内容
    font: str. 英文所用字体
    C_font: str. 中文所用字体
    font_size: int. 字体大小
    font_color: str. 字体颜色(支持'black','red','darkred'，或者自己输入一个RGB列表进去)
    bold: bool. 字体是否加粗

    Return
    ----------
    True
    """
    #清空原本内容
    doc.tables[table].cell(row,col).text = ''
    #添加文字
    if pd.isna(replacement):
        replacement = ''
    run = doc.tables[table].cell(row,col).paragraphs[0].add_run(str(replacement))
    #设置字体和大小
    run.font.name=font
    run.element.rPr.rFonts.set(qn('w:eastAsia'),font)
    run.font.size = Pt(font_size)
    if font_color == 'black':
        color_code = [0,0,0]
    elif font_color == 'darkred':
        color_code = [192,0,0]
    elif font_color == 'red':
        color_code = [255,0,0]
    elif type(font_color)==list and len(font_color)==3:
        color_code = font_color
    else:
        print('不支持该字体颜色')
        color_code = [0,0,0]
    run.font.color.rgb = RGBColor(color_code[0],color_code[1],color_code[2])
    run.bold = bold
    if center:
        doc.tables[table].cell(row,col).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return True



def line_plt(path,data,width=13, height=8,x_interval=6,dpi=300,linewidth=1.5,save=False,fontsize=6):
    """
        Description
        ----------
        画线段图
        
        Parameters
        ----------
        path: str. 保存图片的路径
        data: pd.DataFrame. 每列为需要绘画的数据
        width: int. 图表输出宽度
        height: int. 图表输出高度
        x_interval: int. 横坐标日期间隔N月
        dpi: int. 图表清晰度
        linewidth: int. 图表线段粗细
        save: bool. 是否保存

        Return
        ----------
        None
        """
    fig,ax = plt.subplots(figsize=(width,height),dpi=dpi)
    plt.rcParams['font.sans-serif'] =['Microsoft YaHei']
    plt.rcParams['font.size'] = 5
    line_color = ['#488FD0','#FF9F9F','#BFBFBF','#FFE389']
    if len(data.columns) <= 4:
        for i in range(len(data.columns)):
            ax.plot(data.iloc[:,i],color=line_color[i],label=data.iloc[:,i].name,linewidth=linewidth)
    else:
        for i in range(len(data.columns)):
            ax.plot(data.iloc[:,i],label=data.iloc[:,i].name,linewidth=linewidth)
    figheight = ax.get_window_extent().height
    figwidth = ax.get_window_extent().width
    ax.legend(loc='upper center',
              frameon=False,
              ncol=(int(len(data.columns)/2) if len(data.columns) == 4 else len(data.columns)),
              fontsize=fontsize,
              bbox_to_anchor=(0.5, 1.15))
    ax.spines['right'].set_visible(False)
    ax.spines['top'].set_visible(False)
    ax.spines['left'].set_color('#D9D9D9')
    ax.spines['bottom'].set_color('#D9D9D9')

    ###设置x轴
    #fmt_half_year = mdates.MonthLocator(interval = x_interval)
    #ax.xaxis.set_major_locator(fmt_half_year)
    last_month_start_time=data.index[data.index.map(lambda x:x.strftime('%Y%m'))==data.index[-1].strftime('%Y%m')][0]
    last_month_start_index=data.index.tolist().index(last_month_start_time)
    x_ticks = [data.index[0], data.index[last_month_start_index],
               data.index[int(len(data[:last_month_start_index].index)//4)],
               data.index[int(len(data[:last_month_start_index].index)//2)],
               data.index[int(len(data[:last_month_start_index].index)//4*3)]]
    x_labels = [data.index[0].strftime('%Y-%m'), data.index[last_month_start_index].strftime('%Y-%m'),
                data.index[int(len(data[:last_month_start_index].index)//4)].strftime('%Y-%m'),
                data.index[int(len(data[:last_month_start_index].index)//2)].strftime('%Y-%m'),
                data.index[int(len(data[:last_month_start_index].index)//4*3)].strftime('%Y-%m')]
    ax.set_xticks(x_ticks)
    ax.set_xticklabels(x_labels)

    ###设置y轴
    #_max=data.values.max()+(data.values.max()-data.values.min())/10
    #_min=data.values.min()-(data.values.max()-data.values.min())/10
    #ax.set_ylim(_min,_max)
    #ax.set_ylim(data.values.min()*0.99,data.values.max()*1.01)
    ax.set_ylim(data.values.min()-0.05,data.values.max()+0.1)
    ax.set_xlim(data.index[0],data.index[-1])
    plt.xticks(fontsize=fontsize)
    plt.yticks(fontsize=fontsize)
    plt.tight_layout()
    if save:
        plt.savefig(path+'\线段图.png')
    else:
        plt.show()



def hist_plt(path,data,width,height,dpi=200,barwidth=0.5,save=False,fontsize=8):
    """
        Description
        ----------
        画柱状图
        
        Parameters
        ----------
        path: str. 保存图片的路径
        data: pd.DataFrame. 每列为需要绘画的数据
        width: int. 图表输出宽度
        height: int. 图表输出高度
        x_interval: int. 横坐标日期间隔N月
        dpi: int. 图表清晰度
        linewidth: int. 图表线段粗细
        save: bool. 是否保存

        Return
        ----------
        None
        """
    fig,ax = plt.subplots(figsize=(width,height),dpi=dpi)
    plt.rcParams['font.sans-serif'] =['Microsoft YaHei']
    plt.rcParams['font.size'] = fontsize
    plot = ax.bar(data.index,data.iloc[:,0].values.tolist(),color='#9DC3E6',label=data.index,width=barwidth)
    for value in plot:
        height = value.get_height()
        ax.text(value.get_x() + value.get_width()/2.,
                1.002*height,f'{round(height*100,2)}%', ha='center', va='bottom')
    ax.spines['right'].set_visible(False)
    ax.spines['top'].set_visible(False)
    ax.spines['left'].set_color('#D9D9D9')
    ax.spines['bottom'].set_color('#D9D9D9')
    ax.yaxis.set_major_formatter(ticker.PercentFormatter(xmax=1,decimals=2))
    ax.yaxis.set_major_locator(ticker.FixedLocator([0,0.5,1]))
    ax.set_ylim(0,1)
    plt.xticks(fontsize=fontsize)
    plt.yticks(fontsize=fontsize)
    plt.tight_layout(pad=0.5)
    if save:
        plt.savefig(path+'\柱状图.png')
    else:
        plt.show()



def detect(path=r'输入\精准模板\模板布局'):
    for i in os.listdir(path):
        if os.path.splitext(i)[0]+'.OF' == FundCode:
            return i
    return False 


def find_template():
    #检测基金是否分级
    if_link=type(df_link.loc[FundCode,'link_FundCode'])==list
    if if_link:
        multi = True #True时如果基金有多个份额分类，使用模板二，使交易费率改为一句话。
    else:
        multi = False
    ###读取对应的模板
    if detect():
        doc = docx.Document(path_template+detect())
    else:
        print('不存在精准模板，使用初始模板')
        if language_type=='zh':
            if multi:
                doc = docx.Document(r"输入\中文模板二.docx")
            else:
                doc = docx.Document(r"输入\中文模板.docx")
        else:
           doc = docx.Document(r"输入\英文模板.docx")
    ###信息获取对应的位置
    info_table1=show_tableContent(doc,0)
    info_table2=show_tableContent(doc,-1)
    info_all=info_table1+info_table2
    info = pd.DataFrame(info_all,columns=['','位置']).set_index('') 
    lst_excel_temp=[]#存储merge上位置的excel
    unwriteable_location=[]#存储写不进去的信息
    if_writeable=True#判断是否能写入，一旦有信息写不进去，变量为False，并在结尾处读取初始模板
    #为什么不采用循环，因为每个表都有地方需要做一点处理，不能直接merge
    excel_temp=pd.read_excel(path_result+f'{FundCode}结果汇总.xlsx',sheet_name='基金基本信息',skiprows=[1],index_col=0)
    excel_temp=pd.merge(excel_temp,info,how='left',left_index=True,right_index=True)
    if detect() and pd.isna(excel_temp[excel_temp.index!='基金名称']['位置']).any():
        if_writeable=False
        unwriteable_location+=excel_temp[excel_temp.index!='基金名称'][excel_temp[excel_temp.index!='基金名称']['位置'].isna()].index.tolist()
    else:
        lst_excel_temp.append(excel_temp)

    excel_temp=pd.read_excel(path_result+f'{FundCode}结果汇总.xlsx',sheet_name='业绩排名',skiprows=[1],nrows=5,index_col=0)
    excel_temp=pd.merge(excel_temp,info,how='left',left_index=True,right_index=True)
    if detect() and pd.isna(excel_temp['位置']).any():
        if_writeable=False
        unwriteable_location+=excel_temp[excel_temp['位置'].isna()].index.tolist()
    else:
        lst_excel_temp.append(excel_temp)

    excel_temp=pd.read_excel(path_result+f'{FundCode}结果汇总.xlsx',sheet_name='费率信息',skiprows=[1],index_col=0)
    excel_temp=pd.merge(excel_temp,info,how='left',left_index=True,right_index=True)
    if detect() and pd.isna(excel_temp['位置']).any():
        if_writeable=False
        unwriteable_location+=excel_temp[excel_temp['位置'].isna()].index.tolist()
    else:
        lst_excel_temp.append(excel_temp)

    excel_temp=pd.read_excel(path_result+f'{FundCode}结果汇总.xlsx',sheet_name='前n大股票债券',skiprows=[1],index_col=0,usecols=[0,1,2])
    stock_sort = list('股票'+str(i+1) for i in range(len(excel_temp)))
    excel_temp.index = stock_sort
    excel_temp=pd.merge(excel_temp,info,how='left',left_index=True,right_index=True)
    if detect() and pd.isna(excel_temp['位置']).any() and excel_temp.iloc[0,0]!='--':#有重仓股票且模板写不进去
        if_writeable=False
        unwriteable_location+=excel_temp[excel_temp['位置'].isna()].index.tolist()
    else:
        lst_excel_temp.append(excel_temp)

    excel_temp=pd.read_excel(path_result+f'{FundCode}结果汇总.xlsx',sheet_name='前n大股票债券',skiprows=[1],index_col=0,usecols=[0,3,4])
    stock_sort = list('债券'+str(i+1) for i in range(len(excel_temp)))
    excel_temp.index = stock_sort
    excel_temp=pd.merge(excel_temp,info,how='left',left_index=True,right_index=True)
    if detect() and pd.isna(excel_temp['位置']).any() and excel_temp.iloc[0,0]!='--':#有重仓债券且模板写不进去
        if_writeable=False
        unwriteable_location+=excel_temp[excel_temp['位置'].isna()].index.tolist()
    else:
        lst_excel_temp.append(excel_temp)
    if detect() and if_writeable==True:
        print('全部信息可写入精准模板，使用精准模板')
    elif detect() and if_writeable==False:
        print('部分信息无法写入精准模板，使用初始模板')
        print('以下信息无法写入精准模板：','、'.join(unwriteable_location))
        if language_type=='zh':
            if multi:
                doc = docx.Document(r"输入\中文模板二.docx")
            else:
                doc = docx.Document(r"输入\中文模板.docx")
        else:
           doc = docx.Document(r"输入\英文模板.docx")
        info_table1=show_tableContent(doc,0)
        info_table2=show_tableContent(doc,-1)
        info_all=info_table1+info_table2
        info = pd.DataFrame(info_all,columns=['','位置']).set_index('')
        lst_excel_temp=[]#存储merge上位置的excel
        excel_temp=pd.read_excel(path_result+f'{FundCode}结果汇总.xlsx',sheet_name='基金基本信息',skiprows=[1],index_col=0)
        excel_temp=pd.merge(excel_temp,info,how='left',left_index=True,right_index=True)
        lst_excel_temp.append(excel_temp)
        excel_temp=pd.read_excel(path_result+f'{FundCode}结果汇总.xlsx',sheet_name='业绩排名',skiprows=[1],nrows=5,index_col=0)
        excel_temp=pd.merge(excel_temp,info,how='left',left_index=True,right_index=True)
        lst_excel_temp.append(excel_temp)
        excel_temp=pd.read_excel(path_result+f'{FundCode}结果汇总.xlsx',sheet_name='费率信息',skiprows=[1],index_col=0)
        excel_temp=pd.merge(excel_temp,info,how='left',left_index=True,right_index=True)
        lst_excel_temp.append(excel_temp)
        excel_temp=pd.read_excel(path_result+f'{FundCode}结果汇总.xlsx',sheet_name='前n大股票债券',skiprows=[1],index_col=0,usecols=[0,1,2])
        stock_sort = list('股票'+str(i+1) for i in range(len(excel_temp)))
        excel_temp.index = stock_sort
        excel_temp=pd.merge(excel_temp,info,how='left',left_index=True,right_index=True)
        lst_excel_temp.append(excel_temp)
        excel_temp=pd.read_excel(path_result+f'{FundCode}结果汇总.xlsx',sheet_name='前n大股票债券',skiprows=[1],index_col=0,usecols=[0,3,4])
        stock_sort = list('债券'+str(i+1) for i in range(len(excel_temp)))
        excel_temp.index = stock_sort
        excel_temp=pd.merge(excel_temp,info,how='left',left_index=True,right_index=True)
        lst_excel_temp.append(excel_temp)
    else:
        pass#初始模板一定能写入所有信息，故不再打印提示
    
    return doc,info,lst_excel_temp#lst_excel_temp存储附带位置的excel



def excel_to_word(doc,info,lst_excel_temp,date):
    global today,df_link,FundCode

    #检测基金是否分级
    if_link=type(df_link.loc[FundCode,'link_FundCode'])==list
    if if_link:
        multi = True #True时如果基金有多个份额分类，使用模板二，使交易费率改为一句话。
    else:
        multi = False

    index_detail = pd.read_pickle(path_index)
    try:
        index_detail = index_detail.loc[FundCode,index_detail.loc[FundCode].notnull()]
        index_saved = True
    except:
        index_detail = ''
        index_saved = False
    #0.填充日期
    tables = doc.tables
    table = tables[0]
    for i in range(0,len(table.rows)):
        for j in range(0,len(table.columns)):
            try:
                cell_text=table.cell(i,j).text
            except IndexError:
                continue
            if language_type=='zh':
                if re.search('（日期）+',cell_text):
                    try:
                        new_content = re.search('[\u4e00-\u9fa5]+',table.cell(i,j).text).group()
                        if '净值走势' in new_content:
                            new_content += ' '+today
                            tableText_replace(doc=doc,table=0,row=i,col=j,replacement=new_content,font_size=9,bold=True,center=True)
                        elif '单位净值' in new_content or '累计净值' in new_content:
                            new_content += '（'+today+'）'
                            tableText_replace(doc=doc,table=0,row=i,col=j,replacement=new_content)
                        else:
                            new_content += '（'+date+'）'
                            tableText_replace(doc=doc,table=0,row=i,col=j,replacement=new_content)
                    except IndexError:
                        pass
            else:
                if re.search('(date)+',cell_text):
                    try:
                        new_content = table.cell(i,j).text.split('(')[0]
                        if 'Performance History' in new_content:
                            new_content = new_content+today
                            tableText_replace(doc=doc,table=0,row=i,col=j,replacement=new_content,font_size=9,bold=True,center=True)
                        elif 'NAV' in new_content:
                            new_content = new_content+'('+today+', Yuan)'
                            tableText_replace(doc=doc,table=0,row=i,col=j,replacement=new_content)
                        elif 'Total Net Asset' in new_content:
                            new_content = new_content+'('+date+', Billion Yuan)'
                            tableText_replace(doc=doc,table=0,row=i,col=j,replacement=new_content) 
                        else:
                            new_content = new_content+'('+date+')'
                            tableText_replace(doc=doc,table=0,row=i,col=j,replacement=new_content)
                    except IndexError:
                        pass
    print(f'表一日期改为{date}')
    table = tables[-1]
    for i in range(0,len(table.rows)):
        for j in range(0,len(table.columns)):
            try:
                cell_text=table.cell(i,j).text
            except IndexError:
                continue
            if language_type=='zh':
                if re.search('（日期）+',cell_text):
                    try:
                        new_content = re.search('[\u4e00-\u9fa5]+',table.cell(i,j).text).group()
                        if new_content in ['资产配置','十大重仓股票','五大重仓债券']:
                            new_content += '（'+date+'）'
                            tableText_replace(doc=doc,table=-1,row=i,col=j,replacement=new_content,font_size=9,bold=True,center=True)
                        else:
                            new_content += '（'+date+'）'
                            tableText_replace(doc=doc,table=-1,row=i,col=j,replacement=new_content)
                    except IndexError:
                        pass
            else:
                if re.search('(date)+',cell_text):
                    try:
                        new_content = table.cell(i,j).text.replace('(date)','')
                        if new_content in ['Asset Allocation ','Top 10 Holdings(Stock) ','Top 5 Holdings(Bond) ']:
                            new_content += '('+date+')'
                            tableText_replace(doc=doc,table=-1,row=i,col=j,replacement=new_content,font_size=9,bold=True,center=True)
                    except IndexError:
                        pass
    print(f'表二日期改为{date}')

    #1.基金基本信息
    excel_temp = lst_excel_temp[0].copy()
    fund_name = excel_temp.loc[excel_temp.index=='基金名称','结果'].values[0]
    fund_name = re.search(r'\w+',fund_name).group()
    #写入标题
    if language_type=='zh':
        for num,content in enumerate(doc.paragraphs):
            if '农银' in content.text:
                new_content = excel_temp.loc[excel_temp.index=='基金名称','结果'].values[0][2:]
                run = content.add_run(new_content)
                run.font.name = '微软雅黑'
                run.element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
                run.font.size = Pt(14)
                run.bold = True
                break
    else:
        for num,content in enumerate(doc.paragraphs):
            if 'ABC' in content.text:
                new_content = excel_temp.loc[excel_temp.index=='基金名称','结果'].values[0][3:]
                run = content.add_run(new_content)
                run.font.name = '微软雅黑'
                run.element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
                run.font.size = Pt(14)
                run.bold = True
                break
    #写入基金信息
    for name,df in excel_temp.iterrows():
        if np.sum(pd.notna(df['位置'])):
            if '评级' in name:
                    location = df['位置']
                    tableText_replace(doc=doc,table=0,row=int(location[0]),col=int(location[1]),replacement=df['结果'],font_color='darkred',font_size=6.5,center=True) 
                    continue
            if index_saved:
                if name in index_detail.index.tolist():
                    location = df['位置']
                    font = '微软雅黑'
                    C_font = '微软雅黑'
                    font_size = 8
                    font_color = 'black'
                    bold = False
                    center = False
                    index_temp = index_detail.loc[name]
                    for item in index_temp.items():
                        if item[0] == 'font':
                            font = item[1]
                        elif item[0] == 'c_font':
                            C_font = item[1]
                        elif item[0] == 'font_size':
                            font_size = item[1]
                        elif item[0] == 'font_color':
                            font_color = item[1]
                        elif item[0] == 'bold':
                            bold = item[1]
                        elif item[0] == 'cnter':
                            center = item[1]
                    tableText_replace(doc=doc,table=0,row=int(location[0]),col=int(location[1]),replacement=df['结果'],font=font,C_font=C_font,font_size=font_size,font_color=font_color,bold=bold,center=center)
                    continue
            location = df['位置']
            tableText_replace(doc=doc,table=0,row=int(location[0]),col=int(location[1]),replacement=df['结果'])
    print('基金基本信息替换完成')

    #2.排名信息
    excel_temp = lst_excel_temp[1].copy()
    #需要标红的位置
    se_rows_index=excel_temp['位置'].apply(lambda x:x[0])
    if language_type=='zh':
        ret_index=se_rows_index.loc['本基金收益']
        rank_index=se_rows_index.loc['同类排名']
    else:
        ret_index=se_rows_index.loc['Total Return']
        rank_index=se_rows_index.loc['Rank']
    #写入word
    #本基金收益率大于5%的标红标粗
    fundtype=lst_excel_temp[0].loc['基金类别','结果']
    if fundtype in ['债券型基金','货币型基金','Bond Funds','Money Market Funds']:
        thresh=5
    else:
        thresh=10
    for name,df in excel_temp.iterrows():
        if np.sum(pd.notna(df['位置'])):
            location = list([df['位置'][0],j] for j in df['位置'][1])
            if index_saved:
                font = '微软雅黑'
                C_font = '微软雅黑'
                font_size = 8
                font_color = 'black'
                bold = False
                center = False
                if name in index_detail.index.tolist():
                    index_temp = index_detail.loc[name]
                    for item in index_temp.items():
                        if item[0] == 'font':
                            font = item[1]
                        elif item[0] == 'c_font':
                            C_font = item[1]
                        elif item[0] == 'font_size':
                            font_size = item[1]
                        elif item[0] == 'font_color':
                            font_color = item[1]
                        elif item[0] == 'bold':
                            bold = item[1]
                        elif item[0] == 'cnter':
                            center = item[1]
                for num in range(len(location)):
                    #本基金收益率大于5%的标红标粗
                    if location[num][0]==ret_index:
                        if '%' in df.values[num]:
                            if float(df.values[num][:-1])>=thresh:
                                tableText_replace(doc=doc,table=0,row=int(location[num][0]),col=int(location[num][1]),replacement=df.values[num],font=font,C_font=C_font,font_size=font_size,font_color='darkred',bold=True,center=center)
                                continue
                    #本基金收益率排名在前1/3的标红标粗
                    elif location[num][0]==rank_index:
                        if '/' in df.values[num]:
                            if eval(df.values[num])<=1/3:
                                tableText_replace(doc=doc,table=0,row=int(location[num][0]),col=int(location[num][1]),replacement=df.values[num],font=font,C_font=C_font,font_size=font_size,font_color='darkred',bold=True,center=center)
                                continue
                    tableText_replace(doc=doc,table=0,row=int(location[num][0]),col=int(location[num][1]),replacement=df.values[num],font=font,C_font=C_font,font_size=font_size,font_color=font_color,bold=bold,center=center)
                    continue
                continue
            for num in range(len(location)):
                #本基金收益率大于阈值的标红标粗
                if location[num][0]==ret_index:
                    if '%' in df.values[num]:
                        if float(df.values[num][:-1])>=thresh:
                            tableText_replace(doc=doc,table=0,row=int(location[num][0]),col=int(location[num][1]),replacement=df.values[num],font_size=6.5,bold=True,font_color=[192,0,0],center=True)
                            continue
                #本基金收益率排名在前1/3的标红标粗
                elif location[num][0]==rank_index:
                    if '/' in df.values[num]:
                        if eval(df.values[num])<=1/3:
                            tableText_replace(doc=doc,table=0,row=int(location[num][0]),col=int(location[num][1]),replacement=df.values[num],font_size=6.5,bold=True,font_color=[192,0,0],center=True)
                            continue
                tableText_replace(doc=doc,table=0,row=int(location[num][0]),col=int(location[num][1]),replacement=df.values[num],font_size=6.5,center=True)
    print('基金排名信息替换完成')

    #3.交易费率
    excel_temp = lst_excel_temp[2].copy()
    if multi:
        excel_temp.drop_duplicates(subset=['day','fee','费用大类'],inplace=True)
        def df_to_str(df):
            level = df.name[0]
            fee_type = df.name[1]
            string_list = []
            if fee_type == '费用信息':
                for row in df.iterrows():
                    string_temp = str(row[0])+' '+str(row[1]['fee'])
                    string_list.append(string_temp)
            else:
                for row in df.iterrows():
                    try:
                        if np.isnan(row[1]['day']):
                            string_temp = str(row[1]['fee'])
                        else:
                            string_temp = str(row[1]['day'])+' '+str(row[1]['fee'])
                    except:
                        string_temp = str(row[1]['day'])+' '+str(row[1]['fee'])
                    string_list.append(string_temp)
            final_string = level+'类：'+'；'.join(string_list)
            return final_string
        string_series = excel_temp.groupby(['基金类别','费用大类']).apply(df_to_str)
        def count_characters(text):
            # 初始化统计计数器
            count_english_chinese_semicolon = 0
            count_digits = 0
            count_symbols = 0
            # 遍历字符串中的每个字符
            for char in text:
                if char.isalpha() or char in '：；':
                    count_english_chinese_semicolon += 1
                elif char.isdigit():
                    count_digits += 1
                else:
                    count_symbols += 1
            #return count_english_chinese_semicolon, count_digits, count_symbols
            return count_english_chinese_semicolon+count_digits
        max_len=string_series.apply(count_characters).max()
        string_df = string_series.reset_index()
        def string_combine(df):
            string_list = []
            for row in df.iterrows():
                string_temp = str(row[1][0])
                string_list.append(string_temp)
            final_string = '\n'.join(string_list)
            return final_string
        final_df = pd.DataFrame(string_df.groupby('费用大类').apply(string_combine))
        final_df.columns = ['结果']
        final_df = pd.merge(final_df,info,how='left',left_index=True,right_index=True)
        for name,df in final_df.iterrows():
            if np.sum(pd.notna(df['位置'])):
                location = df['位置']
                if index_saved:
                    font = '微软雅黑'
                    C_font = '微软雅黑'
                    font_size = 8
                    font_color = 'black'
                    bold = False
                    center = False
                    if name in index_detail.index.tolist():
                        index_temp = index_detail.loc[name]
                        for item in index_temp.items():
                            if item[0] == 'font':
                                font = item[1]
                            elif item[0] == 'c_font':
                                C_font = item[1]
                            elif item[0] == 'font_size':
                                font_size = item[1]
                            elif item[0] == 'font_color':
                                font_color = item[1]
                            elif item[0] == 'bold':
                                bold = item[1]
                            elif item[0] == 'cnter':
                                center = item[1]
                    tableText_replace(doc=doc,table=0,row=int(location[0]),col=int(location[1]),replacement=df['结果'],font=font,C_font=C_font,font_size=font_size,font_color=font_color,bold=bold,center=center)
                    continue
                df['结果']=df['结果'].replace('nan ','')
                if max_len<=43:
                    font_size=7.5
                else:
                    font_size=6.5
                tableText_replace(doc=doc,table=0,row=int(location[0]),col=int(location[1]),replacement=df['结果'],font_size=font_size)
        print('基金交易费率替换完成')
    else:
        #excel_temp = excel_temp.loc[excel_temp['基金类别']==FundLevel]
        tables = doc.tables
        table = tables[0]
        #写入收费门槛
        for i in range(0,len(table.rows)):
            for j in range(0,len(table.columns)):
                if re.search('[\u4e00-\u9fa5]+',table.cell(i,j).text):
                    if re.search('[\u4e00-\u9fa5]+',table.cell(i,j).text).group() in ['申购费','赎回费']:
                        try:
                            new_content = excel_temp.loc[excel_temp.index==table.cell(i,j).text,'day'].values[0]
                            if index_saved:
                                font = '微软雅黑'
                                C_font = '微软雅黑'
                                font_size = 8
                                font_color = 'black'
                                bold = False
                                center = False
                                if table.cell(i,j).text in index_detail.index.tolist():
                                    index_temp = index_detail.loc[table.cell(i,j).text]
                                    for item in index_temp.items():
                                        if item[0] == 'font':
                                            font = item[1]
                                        elif item[0] == 'c_font':
                                            C_font = item[1]
                                        elif item[0] == 'font_size':
                                            font_size = item[1]
                                        elif item[0] == 'font_color':
                                            font_color = item[1]
                                        elif item[0] == 'bold':
                                            bold = item[1]
                                        elif item[0] == 'cnter':
                                            center = item[1]
                                tableText_replace(doc=doc,table=0,row=i,col=j,replacement=new_content,font=font,C_font=C_font,font_size=font_size,font_color=font_color,bold=bold,center=center)
                            else:
                                tableText_replace(doc=doc,table=0,row=i,col=j,replacement=new_content)
                        except IndexError:
                            pass
        #写入费率
        for name,df in excel_temp.iterrows():
            if np.sum(pd.notna(df['位置'])):
                location = df['位置']
                if index_saved:
                    font = '微软雅黑'
                    C_font = '微软雅黑'
                    font_size = 8
                    font_color = 'black'
                    bold = False
                    center = False
                    if name in index_detail.index.tolist():
                        index_temp = index_detail.loc[name]
                        for item in index_temp.items():
                            if item[0] == 'font':
                                font = item[1]
                            elif item[0] == 'c_font':
                                C_font = item[1]
                            elif item[0] == 'font_size':
                                font_size = item[1]
                            elif item[0] == 'font_color':
                                font_color = item[1]
                            elif item[0] == 'bold':
                                bold = item[1]
                            elif item[0] == 'cnter':
                                center = item[1]
                    tableText_replace(doc=doc,table=0,row=int(location[0]),col=int(location[1]),replacement=df['fee'],font=font,C_font=C_font,font_size=font_size,font_color=font_color,bold=bold,center=center)
                    continue
                tableText_replace(doc=doc,table=0,row=int(location[0]),col=int(location[1]),replacement=df['fee'])
        #删除多余的收费门槛
        if index_saved==False:
            for i in range(0,len(table.rows)):
                for j in range(0,len(table.columns)):
                    if re.search('[\u4e00-\u9fa5]+',table.cell(i,j).text):
                        if re.search('[\u4e00-\u9fa5]+',table.cell(i,j).text).group() in ['申购费','赎回费']:
                            try:
                                table.cell(i,j).text=''
                            except IndexError:
                                pass
        print('基金交易费率替换完成')

    #4.重仓股票
    excel_temp = lst_excel_temp[3].copy()
    tables = doc.tables
    table = tables[-1]
    for i in range(0,len(table.rows)):
        for j in range(0,len(table.columns)):
            if re.search('[\u4e00-\u9fa5]+',table.cell(i,j).text):
                if re.search('[\u4e00-\u9fa5]+',table.cell(i,j).text).group() in ['股票']:
                    try:
                        new_content = excel_temp.loc[excel_temp.index==table.cell(i,j).text,'股票名称'].values[0]
                        if index_saved:
                            font = '微软雅黑'
                            C_font = '微软雅黑'
                            font_size = 8
                            font_color = 'black'
                            bold = False
                            center = False
                            if table.cell(i,j).text in index_detail.index.tolist():
                                index_temp = index_detail.loc[table.cell(i,j).text]
                                for item in index_temp.items():
                                    if item[0] == 'font':
                                        font = item[1]
                                    elif item[0] == 'c_font':
                                        C_font = item[1]
                                    elif item[0] == 'font_size':
                                        font_size = item[1]
                                    elif item[0] == 'font_color':
                                        font_color = item[1]
                                    elif item[0] == 'bold':
                                        bold = item[1]
                                    elif item[0] == 'cnter':
                                        center = item[1]
                            tableText_replace(doc=doc,table=-1,row=i,col=j,replacement=new_content,font=font,C_font=C_font,font_size=font_size,font_color=font_color,bold=bold,center=True)
                        else:
                            tableText_replace(doc=doc,table=-1,row=i,col=j,replacement=new_content,center=True)
                    except IndexError:
                        pass

    for name,df in excel_temp.iterrows():
        if np.sum(pd.notna(df['位置'])):
            location = df['位置']
            if index_saved:
                font = '微软雅黑'
                C_font = '微软雅黑'
                font_size = 8
                font_color = 'black'
                bold = False
                center = False
                if name in index_detail.index.tolist():
                    index_temp = index_detail.loc[name]
                    for item in index_temp.items():
                        if item[0] == 'font':
                            font = item[1]
                        elif item[0] == 'c_font':
                            C_font = item[1]
                        elif item[0] == 'font_size':
                            font_size = item[1]
                        elif item[0] == 'font_color':
                            font_color = item[1]
                        elif item[0] == 'bold':
                            bold = item[1]
                        elif item[0] == 'cnter':
                            center = item[1]
                tableText_replace(doc=doc,table=-1,row=int(location[0]),col=int(location[1]),replacement=df['股票占净值比'],font=font,C_font=C_font,font_size=font_size,font_color=font_color,bold=bold,center=True)
                continue
            tableText_replace(doc=doc,table=-1,row=int(location[0]),col=int(location[1]),replacement=df['股票占净值比'],center=True)
    print('基金前十大重仓股替换完成')

    #5.重仓债券
    excel_temp = lst_excel_temp[4].copy()
    tables = doc.tables
    table = tables[-1]
    for i in range(0,len(table.rows)):
        for j in range(0,len(table.columns)):
            if re.search('[\u4e00-\u9fa5]+',table.cell(i,j).text):
                if re.search('[\u4e00-\u9fa5]+',table.cell(i,j).text).group() in ['债券']:
                    try:
                        new_content = excel_temp.loc[excel_temp.index==table.cell(i,j).text,'债券名称'].values[0]
                        if index_saved:
                            font = '微软雅黑'
                            C_font = '微软雅黑'
                            font_size = 8
                            font_color = 'black'
                            bold = False
                            center = False
                            if table.cell(i,j).text in index_detail.index.tolist():
                                index_temp = index_detail.loc[table.cell(i,j).text]
                                for item in index_temp.items():
                                    if item[0] == 'font':
                                        font = item[1]
                                    elif item[0] == 'c_font':
                                        C_font = item[1]
                                    elif item[0] == 'font_size':
                                        font_size = item[1]
                                    elif item[0] == 'font_color':
                                        font_color = item[1]
                                    elif item[0] == 'bold':
                                        bold = item[1]
                                    elif item[0] == 'cnter':
                                        center = item[1]
                            tableText_replace(doc=doc,table=-1,row=i,col=j,replacement=new_content,font=font,C_font=C_font,font_size=font_size,font_color=font_color,bold=bold,center=True)
                        else:
                            tableText_replace(doc=doc,table=-1,row=i,col=j,replacement=new_content,center=True)
                    except IndexError:
                        pass
    for name,df in excel_temp.iterrows():
        if np.sum(pd.notna(df['位置'])):
            location = df['位置']
            if index_saved:
                font = '微软雅黑'
                C_font = '微软雅黑'
                font_size = 8
                font_color = 'black'
                bold = False
                center = False
                if name in index_detail.index.tolist():
                    index_temp = index_detail.loc[name]
                    for item in index_temp.items():
                        if item[0] == 'font':
                            font = item[1]
                        elif item[0] == 'c_font':
                            C_font = item[1]
                        elif item[0] == 'font_size':
                            font_size = item[1]
                        elif item[0] == 'font_color':
                            font_color = item[1]
                        elif item[0] == 'bold':
                            bold = item[1]
                        elif item[0] == 'cnter':
                            center = item[1]
                tableText_replace(doc=doc,table=-1,row=int(location[0]),col=int(location[1]),replacement=df['债券占净值比'],font=font,C_font=C_font,font_size=font_size,font_color=font_color,bold=bold,center=True)
                continue
            tableText_replace(doc=doc,table=-1,row=int(location[0]),col=int(location[1]),replacement=df['债券占净值比'],center=True)
    print('基金前五大重仓债券替换完成')

    #6.净值图 资产配置图
    ###写入净值图片
    if multi:
        df_link=pd.read_pickle(path_data+'关联基金.pkl')
        df_link.set_index('FundCode',inplace=True)
        num_link_FundCodes=len(df_link.loc[FundCode,'link_FundCode'])
        excel_temp = lst_excel_temp[0].copy()
        if num_link_FundCodes==1:
            if excel_temp.loc['累计净值','结果']=='--':
                height=4.25
            else:
                height=5.4
        elif num_link_FundCodes==2:
            if excel_temp.loc['累计净值','结果']=='--':
                height=5.15
            else:
                height=7
        width=11.94
    else:
        if language_type=='zh':
            width=11.8
            height=6
        else:
            width=11.8
            height=3.73
    
    data = pd.read_excel(path_result+f'{FundCode}结果汇总.xlsx',sheet_name='净值走势',index_col=0)
    line_plt(r"输出",data,width/2.54,height/2.54,x_interval=15,dpi=300,linewidth=1.5,save=True)
    if_break=False
    flag=0
    for row in doc.tables[0].rows:
        for cell in row.cells:
            if language_type=='zh':
                if '基金规模' in cell.text:
                    flag=1
                if flag==1 and '亿元' in cell.text:
                    flag=2
                if flag==2 and cell.text=='':
                    cell.paragraphs[0].add_run().add_picture(path_result+'线段图.png',Cm(width))
                    if_break=True
                    break
            else:
                if 'Adjusted NAV' in cell.text:
                    flag=1
                if flag==1 and ord(cell.text[0]) in range(48,58):
                    flag=2
                if flag==2 and cell.text=='':
                    cell.paragraphs[0].add_run().add_picture(path_result+'线段图.png',Cm(width))
                    if_break=True
                    break
        if if_break:
            break
    ###写入资产配置图片
    excel_temp=pd.read_excel(path_result+f'{FundCode}结果汇总.xlsx',sheet_name='前n大股票债券',skiprows=[1])
    if excel_temp.iloc[0]['股票名称']=='--' and excel_temp.iloc[0]['债券名称']!='--':
        top_asset_type='bond'
    elif excel_temp.iloc[0]['股票名称']!='--' and excel_temp.iloc[0]['债券名称']=='--':
        top_asset_type='stock'
    else:
        top_asset_type='all'
    data = pd.read_excel(path_result+f'{FundCode}结果汇总.xlsx',sheet_name='资产配置',index_col=0)
    data = data.apply(lambda x:round(x,4))
    if top_asset_type=='all':
        width=6.22
        fontsize=6
    else:
        width=9.34
        fontsize=8
    height=3.10
    hist_plt(r"输出",data,width/2.54,height/2.54,dpi=300,barwidth=0.3,save=True,fontsize=fontsize)
    if_break=False
    flag=0

    if language_type=='zh':
        if top_asset_type=='all' or top_asset_type=='stock':
            tip_ind='十大重仓股票'
        else:
            tip_ind='五大重仓债券'
    else:
        if top_asset_type=='all' or top_asset_type=='stock':
            tip_ind='Top 10 Holdings(Stock)'
        else:
            tip_ind='Top 5 Holdings(Bond)'

    for row in doc.tables[-1].rows:
        for cell in row.cells:
            if tip_ind in cell.text:
                flag=1
            if flag==1 and cell.text=='':
                cell.paragraphs[0].add_run().add_picture(path_result+'柱状图.png',Cm(width))
                if_break=True
                break
        if if_break:
            break
    print('基金净值图和资产配置图写入完成')

    ###写入产品亮点
    if language_type=='zh':
        
        manager=lst_excel_temp[0].loc['基金经理简介','结果'].split('，')[0]
        df_info_newest=pd.read_excel(path_data+'产品信息.xlsx',sheet_name='投资特色、市场展望')
        df_info_newest.set_index('基金经理',inplace=True)
        try:
            TeSe,opinion=df_info_newest.loc[manager]
            tableText_replace(doc=doc,table=-1,row=info.loc['市场展望','位置'][0],col=info.loc['市场展望','位置'][1],replacement=opinion)
            tableText_replace(doc=doc,table=0,row=info.loc['投资特色','位置'][0],col=info.loc['投资特色','位置'][1],replacement=TeSe)
        except:
            print(f'{FundCode}缺少投资特色、市场展望')
        df_info_newest=pd.read_excel(path_data+'产品信息.xlsx',sheet_name='产品亮点',dtype=str)
        df_info_newest['FundCode']=df_info_newest['FundCode'].apply(lambda x:x.zfill(6)+'.OF')
        df_info_newest.set_index('FundCode',inplace=True)
        try:
            LiangDian=df_info_newest.loc[FundCode,'产品亮点']
            tableText_replace(doc=doc,table=0,row=info.loc['产品亮点','位置'][0],col=info.loc['产品亮点','位置'][1],replacement=LiangDian)
        except:
            print(f'{FundCode}缺少产品亮点')
        print('基金产品亮点写入完成')




    doc.save(path_result+f'{fund_name}_{today}.docx')









# %%
### data to excel
#所有参数
path_data=r"./输入/"
path_template=r'./输入/精准模板/模板布局/'
path_index=r'./输入/精准模板/模板参数/完整参数.pkl'
path_result=r"./输出/"
FundCode='001940.OF'
today='2024-03-29'
rptDate_all='2023-12-31'
rptDate_top='2023-12-31'
input_date = '2024Q1'
language_type='zh'



#运行更新excel数据
#读取数据
print('==========1、读取数据==========')
df_link=pd.read_pickle(path_data+'关联基金.pkl')
df_link.set_index('FundCode',inplace=True)

df_FundType=pd.read_excel(path_data+f'上海证券基金类型变更表（{today}）.xlsx',sheet_name='基金类型变更表',dtype='str', header=0, skiprows=1)
df_FundType['基金代码']=df_FundType['基金代码'].apply(lambda x:x+'.OF')
df_FundType.set_index('基金代码',inplace=True)
df_rating=pd.read_excel(path_data+'基金评级查询结果.xlsx',sheet_name='基金评价',dtype='str')
df_rating.set_index('基金代码',inplace=True)
#df_fund_special_date基本不变，故请手动到wind下载
#打开基金数据浏览器
#待选范围为：全部基金(含未成立、已到期)（注：是否可只选择全部基金，尚未测试）
#待选指标为：净值披露首日、基金到期日
df_fund_special_date=pd.read_excel(path_data+f'基金特殊日期（{today}）.xlsx',dtype=str)
df_fund_special_date=df_fund_special_date.iloc[:-2]#删除最后两行的wind水印
df_fund_special_date.rename(columns={'证券代码':'FundCode','证券简称':'FundName'},inplace=True)
df_fund_special_date['基金到期日'].fillna(pd.Timestamp(today)+pd.Timedelta(days=1),inplace=True)
df_fund_special_date['基金到期日']=df_fund_special_date['基金到期日'].apply(pd.Timestamp)
df_fund_special_date=df_fund_special_date[(df_fund_special_date['基金到期日']>=pd.Timestamp(today))]#筛选出尚未到期的基金
df_fund_special_date['净值披露首日']=df_fund_special_date['净值披露首日'].apply(pd.Timestamp)


print('==========2、获取信息==========')
if os.path.exists(path_result+f'{FundCode}结果汇总.xlsx')==False:
    print('不存在结果汇总excel，开始获取信息')
    #运行
    df_info=get_info(df_link,df_rating,FundCode,today,rptDate_all,language_type)
    df_price=get_nav(FundCode,df_info,today,language_type)
    df_asset=get_asset(FundCode,rptDate_top,language_type)
    df_top=get_top(FundCode,rptDate_top)
    df_rank=get_rank(df_FundType,FundCode,today)
    df_fee=get_fee(df_link,FundCode,language_type)
    save(FundCode,df_info,df_price,df_asset,df_top,df_rank,df_fee)
else:
    print('已存在结果汇总excel，跳过信息获取步骤')



### excel to word
print('==========3、信息写入word==========')
doc,info,lst_excel_temp=find_template()
excel_to_word(doc=doc,info=info,lst_excel_temp=lst_excel_temp,date=input_date)








'''

#%%test
doc1 = docx.Document(r"输入\中文模板.docx")
info_table1=show_tableContent(doc1,0)
info_table2=show_tableContent(doc1,-1)
info_all=info_table1+info_table2
list_temp = list(map(lambda x:x[0],info_all))
doc2 = doc = docx.Document(r"输入\中文模板二.docx")
info_table1=show_tableContent(doc2,0)
info_table2=show_tableContent(doc2,-1)
info_all=info_table1+info_table2
for i in info_all:
    if i[0] not in list_temp:
        list_temp.append(i[0])
need_remove = ['农银汇理基金管理有限公司','今年以来','近三个月','近六个月','近一年','近两年','成立以来','交易费率','五大重仓债券','十大重仓股票','占比']
for i in need_remove:
    if i in list_temp:
        list_temp.remove(i)
# %%
a=pd.DataFrame(columns = list_temp)
a.to_pickle(r'输入\精准模板\模板参数\完整参数.pkl')
# %%
a=pd.read_pickle(r'输入\精准模板\模板参数\完整参数.pkl')
# %%
a.loc['660015.OF',['本基金收益','同类排名','同类平均','比较基准','前端申购费率','赎回费率','费用信息']] = [{'font_size':6.5}]
# %%
a.loc['660015.OF',['股票3']] = [{'font_size':10,'font_color':'darkred'}]
# %%
b={'660015.OF':True}
b.to_pickle(r'输入\精准模板\模板参数\Multi.pkl')
# %%
'''
